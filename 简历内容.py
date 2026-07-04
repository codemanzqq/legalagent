# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 读取现有文档
doc = Document(r'D:\系统默认\桌面\汇总\自我成长\简历\5.docx')

# 清空现有内容
doc.paragraphs.clear() if hasattr(doc.paragraphs, 'clear') else None

# 添加标题
def add_heading(text, level=1):
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.bold = True
    runner.font.size = Pt(16 if level == 1 else 12)
    runner.font.name = '微软雅黑'
    runner._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_section_title(text):
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.bold = True
    runner.font.size = Pt(12)
    runner.font.name = '黑体'
    runner._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_content(text, font_size=10.5):
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.font.size = Pt(font_size)
    runner.font.name = '宋体'
    runner._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = Pt(18)
    return p

# 添加简历内容
add_heading('个人简历', level=1)
doc.add_paragraph()  # 空行

# 基本信息
add_section_title('基本信息')
add_content('姓名：[请填写]  |  电话：[请填写]  |  邮箱：[请填写]')
add_content('求职意向：AI 应用开发工程师 / Python 开发工程师 / Java 开发工程师')
doc.add_paragraph()

# 技术栈
add_section_title('技术栈')
add_content('编程语言：Python, Java, JavaScript/TypeScript')
add_content('后端框架：FastAPI, Spring Boot, Spring AI')
add_content('前端框架：Vue 3, React')
add_content('AI 技术：RAG, LangChain, Spring AI Alibaba Graph, LLM 应用开发')
add_content('向量数据库：Milvus, pgvector')
add_content('数据库：MySQL, PostgreSQL, Redis')
add_content('部署运维：Docker, Docker Compose, CI/CD')
doc.add_paragraph()

# 项目经历
add_section_title('项目经历')

# 项目一
add_section_title('项目一：企业级智能数据分析 Agent（DataAgent）')
add_content('项目背景：企业日常数据分析工作依赖专业分析师，业务人员无法直接通过自然语言获取数据洞察，且传统 Text-to-SQL 工具功能单一，无法执行深度分析和生成可视化报告。为降低数据分析门槛，提升企业数据决策效率，基于 Spring AI Alibaba Graph 开发一套企业级智能数据分析 Agent 系统。', font_size=10)
add_content('项目角色：核心开发者', font_size=10)
add_content('技术栈：Spring AI Alibaba Graph、Spring Boot 3.4.8+、Java 17+、Vue 3、Vite、MySQL、向量数据库（pgvector/Milvus 可选）、Python 执行器、ECharts、MCP 协议', font_size=10)
add_content('核心职责：', font_size=10)
add_content('  • 参与项目架构设计与 StateGraph 工作流开发')
add_content('  • 负责核心节点实现（意图识别、SQL 生成与执行、Python 代码执行、报告生成等）')
add_content('  • 主导 RAG 检索增强模块开发，集成向量数据库与混合检索策略')
add_content('  • 实现 Human-in-the-loop 人工反馈机制，支持用户在计划生成阶段干预')
add_content('  • 负责 MCP 服务器集成，支持作为 Tool Server 对接 Claude Desktop 等工具')
add_content('  • 搭建系统性能测试体系，优化并发处理能力')
add_content('项目创新点：', font_size=10)
add_content('  • StateGraph 工作流编排：采用 16+ 节点的状态图工作流模式，实现从用户意图识别到报告生成的全流程自动化')
add_content('  • Python 深度分析与智能报告：内置 Docker/Local Python 执行器，自动生成并执行 Python 代码进行统计分析')
add_content('  • 多模型调度与 MCP 协议支持：内置模型注册表支持运行时动态切换 LLM 与 Embedding 模型')
add_content('项目效果：', font_size=10)
add_content('  • Text-to-SQL 转换准确率达 95%+，Python 代码生成与执行成功率稳定在 90% 以上')
add_content('  • 智能报告生成从小时级缩短至分钟级，提升企业数据决策效率约 70%')
doc.add_paragraph()

# 项目二
add_section_title('项目二：山东省智能政策咨询助手（Policy Agent）')
add_content('项目背景：山东省以旧换新补贴政策涉及品类多、规则复杂，群众咨询量大且问题重复性高，传统客服渠道响应滞后、人力成本高。为提升政策咨询效率与群众满意度，降低客服人力成本，基于 Spring AI 和 RAG 技术开发一套智能政策咨询问答系统。', font_size=10)
add_content('项目角色：后端核心开发者', font_size=10)
add_content('技术栈：Spring Boot 3.4.1、Spring AI 1.0.3、Java 21、React 19、Vite 7、PostgreSQL+pgvector、Redis、MinIO、DashScope（通义千问）、CrossEncoder 重排', font_size=10)
add_content('核心职责：', font_size=10)
add_content('  • 负责 RAG 核心模块开发（文档加载、切片、检索、嵌入模型动态路由）')
add_content('  • 主导 Tool 集成开发（补贴计算工具、发票/文件解析工具、联网搜索工具）')
add_content('  • 实现动态 ChatClient 工厂与多模型管理功能')
add_content('  • 负责会话记忆模块开发，基于 Redis 实现多轮对话上下文管理')
add_content('  • 开发网站爬取与附件提取等多源知识入库管道')
add_content('  • 设计并实现工具失败策略中心统一处理重试与兜底')
add_content('项目创新点：', font_size=10)
add_content('  • RAG 增强与动态嵌入路由：支持文档精细化切片与语义检索，运行时嵌入模型动态路由')
add_content('  • 工具集成与失败策略中心：集成补贴计算、发票解析、联网搜索等工具，统一管理工具失败场景')
add_content('  • 会话事实缓存与多轮对话管理：基于 Redis 实现关键事实结构化写入缓存供多轮对话复用')
add_content('项目效果：', font_size=10)
add_content('  • 政策咨询响应时间从平均 3 分钟缩短至秒级，问答准确率达 95%+，高峰期可承载 500+ 并发咨询')
add_content('  • 补贴计算工具实现自动化核算，准确率 100%，大幅减少人工核算错误')
doc.add_paragraph()

# 项目三
add_section_title('项目三：RAG 智能客服问答系统（Legal_System）')
add_content('项目背景：随着公司业务增长迅速，需要增加客服人力以满足用户对产品的使用咨询、但依旧存在响应不及时等情况，用户满意度有所下降。为提升用户获取信息的效率和满意度，同时减少客服人力，为企业降本增效。基于企业整理的 FAQ 问答对和长文档文件开发一套 RAG 智能客服系统。', font_size=10)
add_content('项目角色：项目负责人', font_size=10)
add_content('技术栈：Qwen、Milvus、混合检索、BM25、RRF、MySQL、Redis、FastAPI', font_size=10)
add_content('核心职责：', font_size=10)
add_content('  • 负责对接产品需求和技术调研')
add_content('  • 主导项目架构设计和代码开发')
add_content('  • 主导项目评估和性能测试体系搭建')
add_content('  • 主导项目迭代优化')
add_content('项目创新点：', font_size=10)
add_content('  • 多级索引与混合检索优化：采用 parent/child chunk 分层分块策略，融合稠密、稀疏向量检索，并引入 BM25 算法')
add_content('  • 精准重排序纠错机制：依托 BM25 和 BGE-M3 完成初步检索，引入 RRF 排序融合算法整合检索分数，再通过 CrossEncoder 模型进行精准重排序')
add_content('项目效果：', font_size=10)
add_content('  • 系统上线后，检索召回率与问答准确率稳定保持在 96% 以上，用户满意度从 86% 提升至 95%')
add_content('  • 大幅降低客服人工成本，整体节省 60% 客服人力')
add_content('  • 业务高峰期可稳定支持 150+ 用户并发咨询，QPM 均值可达 180')
doc.add_paragraph()

# 项目四
add_section_title('项目四：智能食谱生成器（Smart Recipe Planner）')
add_content('项目背景：现代人生活节奏快、饮食不规律，缺乏科学的饮食规划指导；传统食谱应用功能单一，无法根据个人身体数据与饮食偏好提供个性化方案。为帮助用户轻松规划健康饮食，基于 Claude API 开发一套 AI 驱动的个性化食谱生成 Web 应用。', font_size=10)
add_content('项目角色：全栈开发者', font_size=10)
add_content('技术栈：Spring Boot 3.2、Java 17、Spring Security+JWT、React 18、TypeScript、Tailwind CSS、PostgreSQL、Docker、Anthropic SDK（Claude API）', font_size=10)
add_content('核心职责：', font_size=10)
add_content('  • 负责项目架构设计与前后端技术选型')
add_content('  • 主导后端 API 开发（用户认证、食谱生成、聊天助手、数据管理）')
add_content('  • 实现 AI 食谱生成与聊天助手模块，集成 Claude API')
add_content('  • 负责前端页面开发与响应式设计')
add_content('  • 实现 PDF/Word 导出功能')
add_content('  • 配置 Docker 容器化部署与 GitHub Actions CI/CD流水线')
add_content('项目创新点：', font_size=10)
add_content('  • 个性化 AI 食谱生成：基于 Mifflin-St Jeor 公式计算用户每日所需热量，生成 1-30 天可定制的个性化食谱')
add_content('  • AI 聊天助手与上下文感知：支持上下文感知的营养咨询与烹饪指导')
add_content('  • 全栈导出功能：集成 react-pdf 与 docx 库，支持食谱一键导出为 PDF 或 Word 格式')
add_content('项目效果：', font_size=10)
add_content('  • 每天三餐菜品不重复率超 98%，营养均衡度满足中国居民膳食指南标准，用户满意度达 90%+')
add_content('  • AI 聊天助手可准确回答 95%+ 的食谱与营养相关问题')
add_content('  • 约 60% 用户使用过 PDF/Word 导出功能')
doc.add_paragraph()

# 教育背景
add_section_title('教育背景')
add_content('[学校名称]  |  [专业]  |  [学历]  |  [毕业时间]')
doc.add_paragraph()

# 自我评价
add_section_title('自我评价')
add_content('具备 AI 应用开发全流程能力，从需求分析、架构设计到部署上线。熟悉 RAG、Agent 工作流、Text-to-SQL 等主流 AI 应用技术。技术栈横跨 Python(FastAPI) 与 Java(Spring Boot) 生态，前端熟练运用 Vue 3 与 React。有多个从 0 到 1 的项目交付经验，具备良好的技术沟通能力和团队协作精神。')

# 保存文档
output_path = r'D:\系统默认\桌面\汇总\自我成长\简历\5_已修改.docx'
doc.save(output_path)
print(f'简历已保存到：{output_path}')
